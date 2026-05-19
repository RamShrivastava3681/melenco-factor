from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document


ROOT = Path(r"C:\Users\ramsh\Desktop\Whizunik\whiz-factor")
SOURCE = ROOT / "18.07.2024_MODIFI_Framework_Agreement.docx"
TARGET = ROOT / "18.07.2024_WHIZUNIK_Framework_Agreement.docx"


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_runs(paragraph, replacements: dict[str, str]) -> None:
    for run in paragraph.runs:
        updated = run.text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        run.text = updated


def replace_table_text(table, replacements: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_runs(paragraph, replacements)


def main() -> None:
    shutil.copyfile(SOURCE, TARGET)
    doc = Document(TARGET)

    replacements = {
        "MODIFI B.V.": "WHIZUNIK",
        "MODIFI Platform": "WHIZUNIK Platform",
        "MODIFI platform": "WHIZUNIK platform",
        "MODIFI": "WHIZUNIK",
        "Modifi": "Whizunik",
        "Amsterdam": "New Delhi",
        "Netherlands": "India",
        "Dutch Chamber of Commerce": "applicable authorities",
    }

    for paragraph in doc.paragraphs:
        replace_runs(paragraph, replacements)

    for table in doc.tables:
        replace_table_text(table, replacements)

    # Preserve the existing design while correcting the company description.
    set_paragraph_text(
        doc.paragraphs[4],
        "WHIZUNIK, a company incorporated under the laws of India and engaged in trade finance solutions, invoice factoring, export finance, and supply chain financing solutions",
    )
    set_paragraph_text(
        doc.paragraphs[12],
        "The SELLER wishes to make use of the trade finance solutions offered on the WHIZUNIK platform. By entering into this Framework Agreement and subject to the conditions laid down in this Framework Agreement including the accompanying receivables purchase terms and conditions (the RPTC) and any applicable Schedule or Addendum, the SELLER will be admitted as a user of the WHIZUNIK platform.",
    )
    set_paragraph_text(doc.paragraphs[26], "For and behalf of WHIZUNIK")
    set_paragraph_text(doc.paragraphs[30], "New Delhi\tNew Delhi")

    # Clean the visible company label in the fee matrix while keeping the table layout.
    if doc.tables and doc.tables[0].rows and doc.tables[0].rows[0].cells:
        doc.tables[0].rows[0].cells[0].text = "WHIZUNIK"

    props = doc.core_properties
    props.author = "WHIZUNIK"
    props.title = "WHIZUNIK Framework Agreement"
    props.subject = "Framework Agreement"
    props.comments = "Updated from legacy branding to WHIZUNIK while preserving the original layout."

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
